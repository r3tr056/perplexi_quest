import asyncio
import json
import csv
import re
import aiofiles
import logging
import io
from typing import Dict, List, Any, Optional, BinaryIO
from datetime import datetime, timezone
from dataclasses import dataclass
from enum import Enum
import zipfile
import tempfile
import os

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from bibtexparser.bwriter import BibTexWriter
from bibtexparser.bibdatabase import BibDatabase

import httpx
from notion_client import AsyncClient as NotionClient
from notion_client.errors import APIResponseError
import requests

class ExportFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx" 
    HTML = "html"
    MARKDOWN = "markdown"
    JSON = "json"
    CSV = "csv"
    BIBTEX = "bibtex"
    EXCEL = "excel"
    LATEX = "latex"

class IntegrationService(str, Enum):
    NOTION = "notion"
    OBSIDIAN = "obsidian"
    GOOGLE_DOCS = "google_docs"

@dataclass
class ExportConfiguration:
    format: ExportFormat
    include_citations: bool = True
    include_images: bool = True
    include_metadata: bool = True
    include_comments: bool = False
    citation_style: str = "apa"
    template: Optional[str] = None
    custom_settings: Dict[str, Any] = None

@dataclass
class IntegrationConfiguration:
    service: IntegrationService
    credentials: Dict[str, str]
    workspace_id: Optional[str] = None
    folder_id: Optional[str] = None
    custom_settings: Dict[str, Any] = None

logger = logging.getLogger()

class ExportManager:
    def __init__(self):
        self.export_handlers = {
            ExportFormat.PDF: self._export_to_pdf,
            ExportFormat.DOCX: self._export_to_docx,
            ExportFormat.HTML: self._export_to_html,
            ExportFormat.MARKDOWN: self._export_to_markdown,
            ExportFormat.JSON: self._export_to_json,
            ExportFormat.CSV: self._export_to_csv,
            ExportFormat.BIBTEX: self._export_to_bibtex,
            ExportFormat.EXCEL: self._export_to_excel,
            ExportFormat.LATEX: self._export_to_latex
        }
        
        self.integration_handlers = {
            IntegrationService.NOTION: NotionIntegration(),
            IntegrationService.OBSIDIAN: ObsidianIntegration(),
            IntegrationService.GOOGLE_DOCS: GoogleDocsIntegration()
        }

    async def export_research_data(
        self, 
        research_data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            handler = self.export_handlers.get(config.format)
            if not handler:
                raise ValueError(f"Unsupported export format: {config.format}")
            prepared_data = await self._prepare_export_data(research_data, config)
            export_result = await handler(prepared_data, config)
            export_result.update({
                "export_timestamp": datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S'),
                "exported_by": "r3tr056",
                "format": config.format.value,
                "configuration": config.__dict__
            })
            return export_result
        except Exception as e:
            logger.error(f"Export error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def integrate_with_service(
        self, 
        research_data: Dict[str, Any], 
        config: IntegrationConfiguration
    ) -> Dict[str, Any]:
        try:
            handler = self.integration_handlers.get(config.service)
            if not handler:
                raise ValueError(f"Unsupported integration service: {config.service}")
            auth_result = await handler.authenticate(config.credentials)
            if not auth_result["success"]:
                return {"success": False, "error": "Authentication failed"}
            integration_result = await handler.integrate_data(research_data, config)
            integration_result.update({
                "integration_timestamp": datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S'),
                "integrated_by": "r3tr056",
                "service": config.service.value
            })
            return integration_result
        except Exception as e:
            logger.error(f"Integration error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def batch_export(
        self, 
        research_sessions: List[Dict[str, Any]], 
        configurations: List[ExportConfiguration]
    ) -> Dict[str, Any]:
        results = []
        for session in research_sessions:
            session_results = {}
            for config in configurations:
                try:
                    export_result = await self.export_research_data(session, config)
                    session_results[config.format.value] = export_result
                except Exception as e:
                    session_results[config.format.value] = {"success": False, "error": str(e)}
            
            results.append({
                "session_id": session.get("session_id", "unknown"),
                "exports": session_results
            })
        return {
            "batch_export_completed": True,
            "timestamp": datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S'),
            "total_sessions": len(research_sessions),
            "results": results
        }

    async def _prepare_export_data(
        self, 
        research_data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        prepared = {
            "metadata": {
                "title": research_data.get("query", "Research Report"),
                "generated_at": datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S'),
                "generated_by": "r3tr056",
                "research_type": research_data.get("research_type", "standard")
            },
            "content": research_data.get("final_report", {}).get("content", ""),
            "research_plan": research_data.get("plan", {}),
            "research_results": research_data.get("research_results", []),
            "validation_results": research_data.get("validation_results", {}),
            "quality_metrics": research_data.get("quality_metrics", {})
        }
        

        if config.include_citations:
            prepared["citations"] = self._extract_citations_from_results(
                research_data.get("research_results", [])
            )
        if config.include_images:
            prepared["images"] = self._extract_images_from_results(
                research_data.get("research_results", [])
            )
        if config.include_comments:
            prepared["comments"] = research_data.get("comments", {})
        return prepared

    async def _export_to_pdf(
        self, 
        data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1
            )
            story = []
            title = Paragraph(data["metadata"]["title"], title_style)
            story.append(title)
            story.append(Spacer(1, 20))
            if config.include_metadata:
                metadata_data = [
                    ["Generated", data["metadata"]["generated_at"]],
                    ["Research Type", data["metadata"]["research_type"]],
                    ["Generated By", data["metadata"]["generated_by"]]
                ]
                metadata_table = Table(metadata_data)
                metadata_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), '#f0f0f0'),
                    ('TEXTCOLOR', (0, 0), (-1, -1), '#000000'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ]))
                story.append(metadata_table)
                story.append(Spacer(1, 20))

            content_paragraphs = data["content"].split('\n\n')
            for paragraph in content_paragraphs:
                if paragraph.strip():
                    p = Paragraph(paragraph.strip(), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))

            if config.include_citations and data.get("citations"):
                story.append(Spacer(1, 20))
                citations_title = Paragraph("References", styles['Heading2'])
                story.append(citations_title)
                story.append(Spacer(1, 12))
                for citation in data["citations"]:
                    citation_p = Paragraph(citation, styles['Normal'])
                    story.append(citation_p)
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            return {
                "success": True,
                "format": "pdf",
                "content": pdf_bytes,
                "content_type": "application/pdf",
                "filename": f"{data['metadata']['title'].replace(' ', '_')}.pdf",
                "size": len(pdf_bytes)
            }
        except Exception as e:
            logger.error(f"PDF export error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def _export_to_docx(
        self, 
        data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            doc = Document()
            title = doc.add_heading(data["metadata"]["title"], 0)
            title.alignment = 1

            if config.include_metadata:
                doc.add_heading("Document Information", level=2)
                metadata_table = doc.add_table(rows=3, cols=2)
                metadata_table.style = 'Table Grid'
                cells = metadata_table.rows[0].cells
                cells[0].text = "Generated"
                cells[1].text = data["metadata"]["generated_at"]
                cells = metadata_table.rows[1].cells
                cells[0].text = "Research Type"
                cells[1].text = data["metadata"]["research_type"]
                cells = metadata_table.rows[2].cells
                cells[0].text = "Generated By"
                cells[1].text = data["metadata"]["generated_by"]
                doc.add_paragraph()
            
            doc.add_heading("Research Report", level=1)
            content_paragraphs = data["content"].split('\n\n')
            for paragraph in content_paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            if data.get("research_plan"):
                doc.add_page_break()
                doc.add_heading("Research Plan", level=1)
                plan = data["research_plan"]
                if isinstance(plan, dict):
                    for key, value in plan.items():
                        if key != "sub_queries":
                            doc.add_heading(key.replace("_", " ").title(), level=2)
                            doc.add_paragraph(str(value))

            if config.include_citations and data.get("citations"):
                doc.add_page_break()
                doc.add_heading("References", level=1)
                for citation in data["citations"]:
                    p = doc.add_paragraph(citation)
                    p.style = 'Normal'

            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            return {
                "success": True,
                "format": "docx",
                "content": docx_bytes,
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": f"{data['metadata']['title'].replace(' ', '_')}.docx",
                "size": len(docx_bytes)
            }
        except Exception as e:
            logger.error(f"DOCX export error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def _export_to_markdown(
        self, 
        data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            markdown_content = []
            markdown_content.append(f"# {data['metadata']['title']}\n")
            if config.include_metadata:
                markdown_content.append("## Document Information\n")
                markdown_content.append(f"- **Generated:** {data['metadata']['generated_at']}")
                markdown_content.append(f"- **Research Type:** {data['metadata']['research_type']}")
                markdown_content.append(f"- **Generated By:** {data['metadata']['generated_by']}\n")
            markdown_content.append("## Research Report\n")
            markdown_content.append(data["content"])
            markdown_content.append("\n")
            
            if data.get("research_plan"):
                markdown_content.append("## Research Plan\n")
                plan = data["research_plan"]
                if isinstance(plan, dict):
                    for key, value in plan.items():
                        if key != "sub_queries":
                            markdown_content.append(f"### {key.replace('_', ' ').title()}\n")
                            markdown_content.append(f"{value}\n")
            
            if config.include_metadata and data.get("quality_metrics"):
                markdown_content.append("## Quality Metrics\n")
                metrics = data["quality_metrics"]
                for metric, value in metrics.items():
                    if isinstance(value, float):
                        markdown_content.append(f"- **{metric.replace('_', ' ').title()}:** {value:.2f}")
                    else:
                        markdown_content.append(f"- **{metric.replace('_', ' ').title()}:** {value}")
                markdown_content.append("\n")
            
            if config.include_citations and data.get("citations"):
                markdown_content.append("## References\n")
                for i, citation in enumerate(data["citations"], 1):
                    markdown_content.append(f"{i}. {citation}")
                markdown_content.append("\n")
            
            if config.include_images and data.get("images"):
                markdown_content.append("## Images\n")
                for image in data["images"]:
                    if image.get("url"):
                        title = image.get("title", "Research Image")
                        markdown_content.append(f"![{title}]({image['url']})")
                        if image.get("caption"):
                            markdown_content.append(f"\n*{image['caption']}*\n")
                markdown_content.append("\n")
            
            final_markdown = "\n".join(markdown_content)
            return {
                "success": True,
                "format": "markdown",
                "content": final_markdown.encode('utf-8'),
                "content_type": "text/markdown",
                "filename": f"{data['metadata']['title'].replace(' ', '_')}.md",
                "size": len(final_markdown.encode('utf-8'))
            }
        except Exception as e:
            logger.error(f"Markdown export error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def _export_to_json(
        self, 
        data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            json_data = {
                "metadata": data["metadata"],
                "content": data["content"],
                "research_plan": data.get("research_plan", {}),
                "research_results": data.get("research_results", []),
                "validation_results": data.get("validation_results", {}),
                "quality_metrics": data.get("quality_metrics", {})
            }
            if config.include_citations:
                json_data["citations"] = data.get("citations", [])
            if config.include_images:
                json_data["images"] = data.get("images", [])
            if config.include_comments:
                json_data["comments"] = data.get("comments", {})
            json_string = json.dumps(json_data, indent=2, ensure_ascii=False)
            json_bytes = json_string.encode('utf-8')
            return {
                "success": True,
                "format": "json",
                "content": json_bytes,
                "content_type": "application/json",
                "filename": f"{data['metadata']['title'].replace(' ', '_')}.json",
                "size": len(json_bytes)
            }
        except Exception as e:
            logger.error(f"JSON export error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def _export_to_bibtex(
        self, 
        data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            if not data.get("citations"):
                return {"success": False, "error": "No citations available for BibTeX export"}
            bib_db = BibDatabase()
            bib_entries = []
            
            for i, citation in enumerate(data["citations"]):
                entry = {
                    'ENTRYTYPE': 'article',
                    'ID': f'ref{i+1}',
                    'title': citation.get('title', ''),
                    'author': ' and '.join(citation.get('authors', [])),
                    'year': citation.get('publication_date', '')[:4] if citation.get('publication_date') else '',
                    'journal': citation.get('journal', ''),
                    'volume': citation.get('volume', ''),
                    'pages': citation.get('pages', ''),
                    'doi': citation.get('doi', ''),
                    'url': citation.get('url', '')
                }
                entry = {k: v for k, v in entry.items() if v}
                bib_entries.append(entry)
            bib_db.entries = bib_entries
            writer = BibTexWriter()
            bibtex_string = writer.write(bib_db)
            bibtex_bytes = bibtex_string.encode('utf-8')
            return {
                "success": True,
                "format": "bibtex",
                "content": bibtex_bytes,
                "content_type": "application/x-bibtex",
                "filename": f"{data['metadata']['title'].replace(' ', '_')}.bib",
                "size": len(bibtex_bytes)
            }
        except Exception as e:
            logger.error(f"BibTeX export error: {str(e)}")
            return {"success": False, "error": str(e)}

    async def _export_to_latex(
        self, 
        data: Dict[str, Any], 
        config: ExportConfiguration
    ) -> Dict[str, Any]:
        try:
            latex_content = []
            latex_content.extend([
                "\\documentclass[12pt,a4paper]{article}",
                "\\usepackage[utf8]{inputenc}",
                "\\usepackage{amsmath}",
                "\\usepackage{amsfonts}",
                "\\usepackage{amssymb}",
                "\\usepackage{graphicx}",
                "\\usepackage{hyperref}",
                "\\usepackage{cite}",
                "\\usepackage[margin=1in]{geometry}",
                "",
                f"\\title{{{data['metadata']['title']}}}",
                f"\\author{{{data['metadata']['generated_by']}}}",
                f"\\date{{{data['metadata']['generated_at']}}}",
                "",
                "\\begin{document}",
                "\\maketitle",
                ""
            ])
            latex_content.extend([
                "\\begin{abstract}",
                self._escape_latex(data["content"][:500] + "..." if len(data["content"]) > 500 else data["content"]),
                "\\end{abstract}",
                ""
            ])
            latex_content.extend([
                "\\section{Research Report}",
                self._escape_latex(data["content"]),
                ""
            ])
            if data.get("research_plan"):
                latex_content.extend([
                    "\\section{Research Plan}",
                    ""
                ])
                plan = data["research_plan"]
                if isinstance(plan, dict):
                    for key, value in plan.items():
                        if key != "sub_queries":
                            latex_content.extend([
                                f"\\subsection{{{key.replace('_', ' ').title()}}}",
                                self._escape_latex(str(value)),
                                ""
                            ])
            if config.include_citations and data.get("citations"):
                latex_content.extend([
                    "\\section{References}",
                    "\\begin{thebibliography}{99}",
                    ""
                ])
                for i, citation in enumerate(data["citations"], 1):
                    latex_content.append(f"\\bibitem{{ref{i}}} {self._escape_latex(citation)}")
                latex_content.extend([
                    "",
                    "\\end{thebibliography}",
                    ""
                ])
            
            latex_content.append("\\end{document}")
            final_latex = "\n".join(latex_content)
            latex_bytes = final_latex.encode('utf-8')
            return {
                "success": True,
                "format": "latex",
                "content": latex_bytes,
                "content_type": "application/x-latex",
                "filename": f"{data['metadata']['title'].replace(' ', '_')}.tex",
                "size": len(latex_bytes)
            }
        except Exception as e:
            logger.error(f"LaTeX export error: {str(e)}")
            return {"success": False, "error": str(e)}

    def _escape_latex(self, text: str) -> str:
        latex_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '^': '\\textasciicircum{}',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '\\': '\\textbackslash{}'
        }
        for char, escape in latex_chars.items():
            text = text.replace(char, escape)
        return text

    def _extract_citations_from_results(self, research_results: List[Dict[str, Any]]) -> List[str]:
        citations = []
        for result in research_results:
            sources = result.get("sources", [])
            for source in sources:
                citation = f"{source.get('title', 'Unknown Title')}. {source.get('url', '')}"
                if citation not in citations:
                    citations.append(citation)
        return citations

    def _extract_images_from_results(self, research_results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        images = []
        for result in research_results:
            result_images = result.get("images", [])
            for image in result_images:
                if image not in images:
                    images.append(image)
        
        return images

class NotionIntegration:
    async def authenticate(self, credentials: Dict[str, str]) -> Dict[str, Any]:
        try:
            notion = NotionClient(auth=credentials.get("token"))
            await notion.users.retrieve("me")
            return {"success": True}
        except APIResponseError as e:
            logger.error(f"Authentication failed: {e}")
            return {"success": False, "error": str(e)}
        except Exception as e:
            logger.error(f"Unexpected error during authentication: {e}")
            return {"success": False, "error": str(e)}

    async def integrate_data(
        self, 
        research_data: Dict[str, Any], 
        config: IntegrationConfiguration
    ) -> Dict[str, Any]:
        try:
            notion = NotionClient(auth=config.credentials.get("token"))
            page_title = research_data.get("query", "Research Report")
            page_properties = {
                "Name": {
                    "title": [
                        {
                            "text": {
                                "content": page_title
                            }
                        }
                    ]
                }
            }
            
            content_blocks = []
            final_report = research_data.get("final_report", {})
            content = final_report.get("content", "")
            if content:
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs[:10]:
                    paragraph_text = paragraph.strip()
                    if paragraph_text:
                        content_blocks.append({
                            "object": "block",
                            "type": "paragraph",
                            "paragraph": {
                                "rich_text": [
                                    {
                                        "type": "text",
                                        "text": {"content": paragraph_text[:2000]}
                                    }
                                ]
                            }
                        })
            
            page = await notion.pages.create(
                parent={"database_id": config.workspace_id},
                properties=page_properties,
                children=content_blocks
            )
            return {
                "success": True,
                "notion_page_id": page.get("id"),
                "notion_url": page.get("url"),
                "service": "notion"
            }
        except APIResponseError as e:
            logger.error(f"Notion API error: {str(e)}")
            return {"success": False, "error": str(e)}
        except Exception as e:
            logger.error(f"Notion integration error: {str(e)}")
            return {"success": False, "error": str(e)}

class ObsidianIntegration:
    async def authenticate(self, credentials: Dict[str, str]) -> Dict[str, Any]:
        vault_path = credentials.get("vault_path")
        if vault_path and os.path.exists(vault_path):
            return {"success": True}
        return {"success": False, "error": "Invalid vault path"}

    async def integrate_data(
        self, 
        research_data: Dict[str, Any], 
        config: IntegrationConfiguration
    ) -> Dict[str, Any]:
        try:
            vault_path = config.credentials.get("vault_path")
            if not vault_path or not os.path.exists(vault_path):
                return {"success": False, "error": "Invalid vault path"}
            
            markdown_content = []
            title = research_data.get("query", "Research Report")
            markdown_content.append(f"# {title}\n")
            markdown_content.append("## Metadata")
            markdown_content.append(f"- Created: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}")
            markdown_content.append(f"- Type: {research_data.get('research_type', 'standard')}")
            markdown_content.append(f"- Generated by: r3tr056\n")
            
            tags = ["#research", "#perplexiquest"]
            if research_data.get("domain"):
                tags.append(f"#{research_data['domain']}")
            markdown_content.append("Tags: " + " ".join(tags) + "\n")
            
            if research_data.get("final_report", {}).get("content"):
                markdown_content.append("## Research Report\n")
                markdown_content.append(research_data["final_report"]["content"])
                markdown_content.append("\n")
            
            if research_data.get("plan"):
                markdown_content.append("## Research Plan\n")
                plan = research_data["plan"]
                if isinstance(plan, dict):
                    for key, value in plan.items():
                        markdown_content.append(f"### {key.replace('_', ' ').title()}\n")
                        markdown_content.append(f"{value}\n")
            
            if research_data.get("research_results"):
                markdown_content.append("## Sources\n")
                for result in research_data["research_results"]:
                    sources = result.get("sources", [])
                    for source in sources[:5]:
                        title = source.get("title", "Unknown")
                        url = source.get("url", "")
                        markdown_content.append(f"- [{title}]({url})")
                markdown_content.append("\n")

            def filename_san(title: str) -> str:
                return re.sub(r'[\\/*?:"<>|]', "_", title)
            
            filename = f"{filename_san(title)}.md"
            file_path = os.path.join(vault_path, filename)
            
            async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                await f.write('\n'.join(markdown_content))
            
            return {
                "success": True,
                "obsidian_file_path": file_path,
                "filename": filename,
                "service": "obsidian"
            }
        except Exception as e:
            logger.error(f"Obsidian integration error: {str(e)}")
            return {"success": False, "error": str(e)}

class GoogleDocsIntegration:
    SCOPES = ['https://www.googleapis.com/auth/documents']
    def __init__(self):
        self.creds = None
        self.service = None

    async def authenticate(self, credentials):
        try:
            if os.path.exists('token.json'):
                self.creds = Credentials.from_authorized_user_file('token.json', self.SCOPES)
            if not self.creds or not self.creds.valid:
                if self.creds and self.creds.expired and self.creds.refresh_token:
                    self.creds.refresh(Request())
                else:
                    flow = InstalledAppFlow.from_client_secrets_file('credentials.json', self.SCOPES)
                    self.creds = flow.run_local_server(port=0)
                with open('token.json', 'w') as token:
                    token.write(self.creds.to_json())
            self.service = build('docs', 'v1', credentials=self.creds)
            return {"success": True}
        except Exception as e:
            logger.error(f"Authentication failed: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    async def integrate_data(self, data: Dict[str, Any], config: Any) -> Dict[str, Any]:
        try:
            if not self.service:
                return {"success": False, "error": "Service not initialized. Call authenticate() first."}
            
            title = data.get("query", "Research Report")
            document = self.service.documents().create(body={'title': title}).execute()
            document_id = document.get('documentId')
            logger.info(f"Created document with ID: {document_id}")
            content = self._prepare_content(data)
            requests = [
                {
                    'insertText': {
                        'location': {
                            'index': 1,
                        },
                        'text': content
                    }
                }
            ]
            self.service.documents().batchUpdate(documentId=document_id, body={'requests': requests}).execute()
            doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
            return {
                "success": True,
                "google_docs_url": doc_url,
                "google_docs_id": document_id,
                "service": "google_docs"
            }
        except HttpError as e:
            logger.error(f"An HTTP error occured: {e}", exc_info=True)
            return {"success": False, "error": str(e)}
        except Exception as e:
            logger.error(f"An error occured: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _prepare_content(self, research_data: Dict[str, Any]) -> str:
        lines = []
        title = research_data.get("query", "Research Report")
        lines.append(f"{title}\n")
        lines.append("Metadata")
        lines.append(f"- Created: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"- Type: {research_data.get('research_type', 'standard')}")
        lines.append(f"- Generated by: r3tr056\n")

        tags = ["#research", "#perplexiquest"]
        if research_data.get("domain"):
            tags.append(f"#{research_data['domain']}")
        lines.append("Tags: " + " ".join(tags) + "\n")

        final_report = research_data.get("final_report", {}).get("content")
        if final_report:
            lines.append("Research Report\n")
            lines.append(final_report + "\n")

        plan = research_data.get("plan")
        if plan and isinstance(plan, dict):
            lines.append("Research Plan\n")
            for key, value in plan.items():
                lines.append(f"{key.replace('_', ' ').title()}\n")
                lines.append(f"{value}\n")

        research_results = research_data.get("research_results")
        if research_results:
            lines.append("Sources\n")
            for result in research_results:
                sources = result.get("sources", [])
                for source in sources[:5]:
                    source_title = source.get("title", "Unknown")
                    url = source.get("url", "")
                    lines.append(f"- {source_title} ({url})")
            lines.append("\n")

        return '\n'.join(lines)
